from __future__ import annotations

from collections import defaultdict
from datetime import UTC, datetime
import io
import json
import sqlite3
from typing import Any

from docx import Document

from mckainsey.config import Settings
from mckainsey.services.config_service import ConfigService
from mckainsey.services.llm_client import GeminiChatClient
from mckainsey.services.memory_service import MemoryService
from mckainsey.services.storage import SimulationStore


def _clean_report_text(value: Any) -> str:
    text = str(value or "").replace("**", "").replace("`", "").strip()
    return " ".join(text.split())


def _format_metric_value(value: float, unit: str) -> str:
    if unit == "%":
        return f"{value:.1f}%"
    if unit == "/10":
        return f"{value:.1f}/10"
    return f"{value:.1f}"


class ReportService:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.store = SimulationStore(settings.simulation_db_path)
        self.llm = GeminiChatClient(settings)
        self.memory = MemoryService(settings)

    def generate_structured_report(self, simulation_id: str, use_case: str | None = None) -> dict[str, Any]:
        existing = self.store.get_report_state(simulation_id)
        if existing and existing.get("status") == "completed":
            return existing

        agents = self.store.get_agents(simulation_id)
        interactions = self.store.get_interactions(simulation_id)
        if not agents:
            raise ValueError(f"Simulation not found: {simulation_id}")

        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}
        population = self.store.get_population_artifact(simulation_id) or {}
        baseline = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="baseline")
        final = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="final")
        events = self.store.list_simulation_events(simulation_id)

        payload: dict[str, Any] = {}
        if self._should_request_structured_report_seed():
            prompt = self._build_structured_report_prompt(
                simulation_id=simulation_id,
                use_case=use_case,
                knowledge=knowledge,
                population=population,
                agents=agents,
                interactions=interactions,
                baseline=baseline,
                final=final,
                events=events,
            )
            try:
                raw = self.llm.complete_required(
                    prompt,
                    system_prompt=(
                        "You are McKAInsey ReportAgent. Return valid JSON only using the requested schema. "
                        "Every claim must be grounded in provided evidence."
                    ),
                )
            except Exception:  # noqa: BLE001
                raw = ""
            try:
                parsed_payload = _parse_json_object(raw) if raw else {}
            except Exception:  # noqa: BLE001
                parsed_payload = {}
            if isinstance(parsed_payload, dict):
                payload = parsed_payload

        normalized = self._normalize_structured_report_payload(simulation_id, payload)
        return self._enrich_structured_report_payload(
            simulation_id,
            normalized,
            use_case=use_case,
            agents=agents,
            interactions=interactions,
            baseline=baseline,
            final=final,
            events=events,
            knowledge=knowledge,
            population=population,
        )

    def _should_request_structured_report_seed(self) -> bool:
        # Local Ollama models have been the slowest and least reliable source of
        # large JSON objects in live mode. We still build a report from the real
        # simulation artifacts, but skip the heavyweight seed request here.
        return self.llm.provider != "ollama"

    def build_report(self, simulation_id: str) -> dict[str, Any]:
        cached = self.store.get_cached_report(simulation_id)
        if cached:
            return cached

        agents = self.store.get_agents(simulation_id)
        interactions = self.store.get_interactions(simulation_id)
        if not agents:
            raise ValueError(f"Simulation not found: {simulation_id}")

        pre = [float(a["opinion_pre"]) for a in agents]
        post = [float(a["opinion_post"]) for a in agents]

        by_area_pre: dict[str, list[float]] = defaultdict(list)
        by_area_post: dict[str, list[float]] = defaultdict(list)
        by_income_post: dict[str, list[float]] = defaultdict(list)
        influence: dict[str, float] = defaultdict(float)
        agent_persona: dict[str, dict[str, Any]] = {}
        for a in agents:
            area = str(a["persona"].get("planning_area", "Unknown"))
            income = str(a["persona"].get("income_bracket", "Unknown"))
            by_area_pre[area].append(float(a["opinion_pre"]))
            by_area_post[area].append(float(a["opinion_post"]))
            by_income_post[income].append(float(a["opinion_post"]))
            agent_persona[a["agent_id"]] = a["persona"]

        for i in interactions:
            if i.get("target_agent_id"):
                influence[i["actor_agent_id"]] += abs(float(i.get("delta", 0)))

        last_reason_by_agent: dict[str, str] = {}
        for i in interactions:
            text = str(i.get("content", "")).strip()
            if text:
                last_reason_by_agent[i["actor_agent_id"]] = text[:240]

        influential_agents: list[dict[str, Any]] = []
        for agent_id, score in sorted(influence.items(), key=lambda x: x[1], reverse=True)[:10]:
            persona = agent_persona.get(agent_id, {})
            influential_agents.append(
                {
                    "agent_id": agent_id,
                    "influence_score": round(score, 4),
                    "planning_area": str(persona.get("planning_area", "Unknown")),
                    "occupation": str(persona.get("occupation", "Unknown")),
                    "income_bracket": str(persona.get("income_bracket", "Unknown")),
                    "latest_argument": last_reason_by_agent.get(agent_id, "No recent argument captured."),
                }
            )

        area_metrics: list[dict[str, Any]] = []
        for area, post_scores in by_area_post.items():
            pre_scores = by_area_pre.get(area, [])
            post_mean = _mean(post_scores)
            pre_mean = _mean(pre_scores)
            approval_post = _approval(post_scores)
            mean_shift = post_mean - pre_mean
            friction = abs(mean_shift) * (1 - approval_post)
            area_metrics.append(
                {
                    "planning_area": area,
                    "avg_pre_opinion": round(pre_mean, 4),
                    "avg_post_opinion": round(post_mean, 4),
                    "approval_post": round(approval_post, 4),
                    "mean_shift": round(mean_shift, 4),
                    "friction_index": round(friction, 4),
                    "cohort_size": len(post_scores),
                }
            )

        top_dissenting = sorted(
            area_metrics,
            key=lambda x: (x["approval_post"], -x["friction_index"]),
        )[:8]

        income_metrics = [
            {
                "income_bracket": income,
                "approval_post": round(_approval(scores), 4),
                "avg_post_opinion": round(_mean(scores), 4),
                "cohort_size": len(scores),
            }
            for income, scores in by_income_post.items()
        ]

        arguments_for = [
            {
                "text": i.get("content", ""),
                "agent_id": i["actor_agent_id"],
                "round_no": i["round_no"],
                "strength": round(abs(float(i.get("delta", 0))), 4),
            }
            for i in interactions
            if float(i.get("delta", 0)) > 0
        ]
        arguments_for = sorted(arguments_for, key=lambda x: x["strength"], reverse=True)[:12]

        arguments_against = [
            {
                "text": i.get("content", ""),
                "agent_id": i["actor_agent_id"],
                "round_no": i["round_no"],
                "strength": round(abs(float(i.get("delta", 0))), 4),
            }
            for i in interactions
            if float(i.get("delta", 0)) < 0
        ]
        arguments_against = sorted(arguments_against, key=lambda x: x["strength"], reverse=True)[:12]

        executive_summary = self.llm.complete(
            prompt=(
                f"Generate a concise executive summary for simulation {simulation_id}. "
                f"Pre approval={_approval(pre):.2f}, post approval={_approval(post):.2f}, "
                f"net shift={_mean(post)-_mean(pre):.2f}. "
                f"Top dissent cohorts={top_dissenting[:3]}."
            ),
            system_prompt="You are ReportAgent. Return concise strategic summary.",
        )

        recommendations = self._recommend(
            simulation_id=simulation_id,
            top_dissenting=top_dissenting,
            income_metrics=income_metrics,
            arguments_for=arguments_for,
            arguments_against=arguments_against,
        )

        report = {
            "simulation_id": simulation_id,
            "executive_summary": executive_summary,
            "approval_rates": {
                "stage3a": round(_approval(pre), 4),
                "stage3b": round(_approval(post), 4),
                "delta": round(_approval(post) - _approval(pre), 4),
            },
            "top_dissenting_demographics": top_dissenting,
            "friction_by_planning_area": sorted(area_metrics, key=lambda x: x["friction_index"], reverse=True),
            "income_cohorts": sorted(income_metrics, key=lambda x: x["approval_post"]),
            "influential_agents": influential_agents,
            "key_arguments_for": arguments_for,
            "key_arguments_against": arguments_against,
            "recommendations": recommendations,
        }

        self.store.cache_report(simulation_id, report)
        return report

    def report_chat(self, simulation_id: str, message: str) -> str:
        report = self.build_report(simulation_id)
        prompt = (
            f"Report JSON:\n{report}\n\n"
            f"User asks: {message}\n"
            "Provide a direct, data-grounded answer with concrete cohort references."
        )
        return self.llm.complete_required(prompt, system_prompt="You are McKAInsey ReportAgent.")

    def report_chat_payload(self, simulation_id: str, message: str) -> dict[str, Any]:
        report = self.build_report(simulation_id)
        zep_context = self.memory.search_simulation_context(simulation_id, message, limit=8)
        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}
        zep_excerpt = "\n".join(
            f"- {item['content']}"
            for item in zep_context["episodes"][:6]
        )
        knowledge_excerpt = "\n".join(self._knowledge_context_lines(knowledge))
        prompt = (
            f"Report JSON:\n{report}\n\n"
            f"Original document context:\n{knowledge_excerpt or '- none'}\n\n"
            f"Relevant Zep Cloud memory search results:\n{zep_excerpt or '- none'}\n\n"
            f"User asks: {message}\n"
            "Provide a direct, data-grounded answer with concrete cohort references."
        )
        response = self.llm.complete_required(prompt, system_prompt="You are McKAInsey ReportAgent.")
        return {
            "session_id": simulation_id,
            "simulation_id": simulation_id,
            "response": response,
            "model_provider": self.llm.provider,
            "model_name": self.llm.model_name,
            "gemini_model": self.llm.model_name,
            "zep_context_used": zep_context["zep_context_used"],
        }

    def build_v2_report(self, simulation_id: str, use_case: str | None = None) -> dict[str, Any]:
        from mckainsey.services.metrics_service import MetricsService

        agents = self.store.get_agents(simulation_id)
        interactions = self.store.get_interactions(simulation_id)
        if not agents:
            raise ValueError(f"Simulation not found: {simulation_id}")
        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}

        config_service = ConfigService(self.settings)
        metrics_service = MetricsService(config_service)

        # Resolve analysis questions and metadata
        analysis_questions = self._resolve_analysis_questions(simulation_id=simulation_id, use_case=use_case)
        insight_block_configs = self._resolve_insight_blocks(use_case)
        preset_section_configs = self._resolve_preset_sections(use_case)

        simulation = self.store.get_simulation(simulation_id) or {}
        stored_round_count = int(simulation.get("rounds", 0) or 0)
        interaction_round_count = max((int(item.get("round_no", 0) or 0) for item in interactions), default=0)
        round_count = stored_round_count or interaction_round_count
        evidence_pool = self._extract_evidence(interactions, agents)

        # ── Metric deltas for quantitative questions ──
        baseline_records = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="baseline")
        final_records = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="final")
        metric_deltas: list[dict[str, Any]] = []
        for q in analysis_questions:
            if q.get("type") == "open-ended":
                continue
            name = q.get("metric_name", "")
            if not name:
                continue
            # Compute R1 and final values
            r1_agents = self._agents_from_checkpoint(baseline_records) or agents
            final_agents = self._agents_from_checkpoint(final_records) or agents
            r1_val = self._compute_metric_value(q, r1_agents)
            final_val = self._compute_metric_value(q, final_agents)
            metric_unit = "%" if q.get("type") == "yes-no" or "threshold" in q or str(q.get("metric_unit", "")).strip() == "%" else str(q.get("metric_unit", "/10") or "/10")
            initial_display = _format_metric_value(r1_val, metric_unit)
            final_display = _format_metric_value(final_val, metric_unit)
            metric_deltas.append({
                "metric_name": name,
                "metric_label": q.get("metric_label", name),
                "metric_unit": metric_unit,
                "initial_value": r1_val,
                "final_value": final_val,
                "delta": round(final_val - r1_val, 2),
                "direction": "up" if final_val > r1_val else ("down" if final_val < r1_val else "flat"),
                "report_title": q.get("report_title", name),
                "initial_display": initial_display,
                "final_display": final_display,
                "delta_display": f"{initial_display} -> {final_display}",
            })

        # ── Build report sections from analysis questions ──
        sections: list[dict[str, Any]] = []
        for q in analysis_questions:
            question_text = q.get("question", "")
            q_type = q.get("type", "scale")
            answer = self._answer_guiding_question(simulation_id, question_text, agents, interactions)
            answer = _clean_report_text(answer)

            section: dict[str, Any] = {
                "question": question_text,
                "report_title": q.get("report_title", question_text[:60]),
                "type": q_type,
                "answer": answer,
                "evidence": evidence_pool[:3],
            }

            # Add metric spotlight for quantitative questions
            if q_type != "open-ended":
                delta_entry = next((d for d in metric_deltas if d["metric_name"] == q.get("metric_name")), None)
                if delta_entry:
                    section["metric"] = delta_entry

            sections.append(section)

        # ── Compute insight blocks ──
        insight_blocks: list[dict[str, Any]] = []
        for block_cfg in insight_block_configs:
            block_type = block_cfg.get("type", "")
            result = metrics_service.compute_insight_block(
                block_type=block_type,
                agents=agents,
                interactions=interactions,
                analysis_questions=analysis_questions,
                metric_ref=block_cfg.get("metric_ref"),
                count=block_cfg.get("count", 5),
            )
            insight_blocks.append({
                "type": block_type,
                "title": block_cfg.get("title", block_type),
                "description": block_cfg.get("description", ""),
                "data": result,
            })

        # ── Generate preset sections via LLM ──
        preset_sections: list[dict[str, Any]] = []
        for preset in preset_section_configs:
            title = preset.get("title", "")
            prompt = preset.get("prompt", "")
            if prompt:
                answer = self._answer_guiding_question(simulation_id, prompt, agents, interactions)
            else:
                answer = ""
            preset_sections.append({"title": title, "answer": answer})

        # ── Executive summary ──
        executive_summary = self._build_v2_executive_summary_from_metrics(
            simulation_id=simulation_id,
            metric_deltas=metric_deltas,
            round_count=round_count,
            agent_count=len(agents),
        )

        return {
            "session_id": simulation_id,
            "generated_at": datetime.now(UTC).isoformat(),
            "executive_summary": _clean_report_text(executive_summary),
            "metric_deltas": metric_deltas,
            "quick_stats": {
                "agent_count": len(agents),
                "round_count": round_count,
                "model": self.llm.model_name,
                "provider": self.llm.provider,
            },
            "sections": sections,
            "insight_blocks": insight_blocks,
            "preset_sections": preset_sections,
        }

    def export_v2_report_docx(self, simulation_id: str, report: dict[str, Any] | None = None, use_case: str | None = None) -> bytes:
        payload = report or self.build_v2_report(simulation_id, use_case=use_case)
        document = Document()
        document.add_heading("McKAInsey Analysis Report", level=0)
        document.add_paragraph(f"Session: {payload.get('session_id', simulation_id)}")
        document.add_paragraph(f"Generated: {payload.get('generated_at', '')}")

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(str(payload.get("executive_summary", "")))

        quick_stats = payload.get("quick_stats", {})
        if isinstance(quick_stats, dict):
            document.add_heading("Quick Stats", level=1)
            document.add_paragraph(
                f"Agents: {quick_stats.get('agent_count', 0)} | Rounds: {quick_stats.get('round_count', 0)}"
            )
            document.add_paragraph(
                f"Model: {quick_stats.get('model', '')} ({quick_stats.get('provider', '')})"
            )

        metric_deltas = payload.get("metric_deltas", [])
        if isinstance(metric_deltas, list) and metric_deltas:
            document.add_heading("Metric Deltas", level=1)
            for metric in metric_deltas:
                if not isinstance(metric, dict):
                    continue
                label = str(metric.get("metric_label", metric.get("metric_name", "Metric")))
                initial_value = metric.get("initial_value", 0)
                final_value = metric.get("final_value", 0)
                delta = metric.get("delta", 0)
                unit = str(metric.get("metric_unit", ""))
                document.add_paragraph(
                    f"{label}: {initial_value}{unit} -> {final_value}{unit} ({delta:+}{unit})",
                    style="List Bullet",
                )

        document.add_heading("Analysis Question Sections", level=1)
        for section in payload.get("sections", []):
            if not isinstance(section, dict):
                continue
            section_title = str(section.get("report_title") or section.get("question") or "Section")
            document.add_heading(section_title, level=2)
            if section.get("question"):
                document.add_paragraph(f"Question: {section.get('question')}")
            document.add_paragraph(str(section.get("answer", "")))
            evidence = section.get("evidence", [])
            if isinstance(evidence, list) and evidence:
                document.add_paragraph("Evidence:")
                for item in evidence:
                    if isinstance(item, dict):
                        quote = str(item.get("quote", "")).strip()
                        agent_id = str(item.get("agent_id", ""))
                        post_id = str(item.get("post_id", ""))
                        document.add_paragraph(
                            f"{agent_id} / {post_id}: {quote}",
                            style="List Bullet",
                        )

        insight_blocks = payload.get("insight_blocks", [])
        if isinstance(insight_blocks, list) and insight_blocks:
            document.add_heading("Use-Case Insights", level=1)
            for block in insight_blocks:
                if not isinstance(block, dict):
                    continue
                document.add_heading(str(block.get("title", "Insight")), level=2)
                description = str(block.get("description", "")).strip()
                if description:
                    document.add_paragraph(description)

        preset_sections = payload.get("preset_sections", [])
        if isinstance(preset_sections, list) and preset_sections:
            document.add_heading("Preset Sections", level=1)
            for section in preset_sections:
                if not isinstance(section, dict):
                    continue
                document.add_heading(str(section.get("title", "Section")), level=2)
                document.add_paragraph(str(section.get("answer", "")))

        document.add_paragraph(
            f"Methodology: Simulated agents={quick_stats.get('agent_count', 0)}, "
            f"rounds={quick_stats.get('round_count', 0)}, "
            f"model={quick_stats.get('model', '')} ({quick_stats.get('provider', '')})."
        )

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _resolve_guiding_questions(self, use_case: str | None) -> list[str]:
        if use_case:
            config_service = ConfigService(self.settings)
            try:
                checkpoint_questions = config_service.get_checkpoint_questions(use_case)
            except Exception:  # noqa: BLE001
                checkpoint_questions = []
            checkpoint_prompts = [
                str(item.get("question", "")).strip()
                for item in checkpoint_questions
                if isinstance(item, dict) and str(item.get("question", "")).strip()
            ]
            if checkpoint_prompts:
                return checkpoint_prompts
            try:
                sections = config_service.get_report_sections(use_case)
            except Exception:  # noqa: BLE001
                sections = []
            report_prompts = [
                str(item.get("prompt") or item.get("title") or "").strip()
                for item in sections
                if isinstance(item, dict) and str(item.get("prompt") or item.get("title") or "").strip()
            ]
            if report_prompts:
                return report_prompts

        return [
            "What are the major shifts in opinion across rounds?",
            "Which arguments most strongly support the policy?",
            "Which arguments most strongly oppose the policy?",
        ]

    def _extract_evidence(self, interactions: list[dict[str, Any]], agents: list[dict[str, Any]]) -> list[dict[str, Any]]:
        agent_names: dict[str, str] = {}
        for agent in agents:
            agent_id = str(agent.get("agent_id") or agent.get("id") or "").strip()
            if not agent_id:
                continue
            display_name = _clean_report_text(
                agent.get("name")
                or agent.get("agent_name")
                or agent.get("display_name")
                or agent.get("label")
                or agent_id
            )
            agent_names[agent_id] = display_name or agent_id

        evidence: list[dict[str, Any]] = []
        for row in interactions:
            quote = str(row.get("content", "")).strip()
            if not quote:
                continue
            agent_id = str(row.get("actor_agent_id", "")).strip()
            evidence.append(
                {
                    "agent_id": agent_id,
                    "agent_name": agent_names.get(agent_id, agent_id),
                    "post_id": str(row.get("post_id") or row.get("id") or ""),
                    "quote": quote[:280],
                }
            )
        return evidence

    def _knowledge_context_lines(self, knowledge: dict[str, Any]) -> list[str]:
        if not isinstance(knowledge, dict):
            return []

        lines: list[str] = []
        summary = _clean_report_text(knowledge.get("summary"))
        if summary:
            lines.append(f"Document summary: {summary}")

        document = knowledge.get("document")
        if isinstance(document, dict):
            source_path = _clean_report_text(document.get("source_path"))
            if source_path:
                lines.append(f"Source document: {source_path}")
            text_length = document.get("text_length")
            if text_length is not None:
                lines.append(f"Document length: {text_length}")
            sources = document.get("sources")
            if isinstance(sources, list):
                for source in sources[:3]:
                    if not isinstance(source, dict):
                        continue
                    source_path = _clean_report_text(source.get("source_path"))
                    if source_path:
                        lines.append(f"Document source: {source_path}")
        return lines

    def _answer_guiding_question(
        self,
        simulation_id: str,
        question: str,
        agents: list[dict[str, Any]],
        interactions: list[dict[str, Any]],
    ) -> str:
        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}
        knowledge_context = "\n".join(self._knowledge_context_lines(knowledge))
        prompt = (
            f"Simulation ID: {simulation_id}\n"
            f"Guiding question: {question}\n"
            f"Original document context:\n{knowledge_context or '- none'}\n"
            f"Agent sample size: {len(agents)}\n"
            f"Recent interactions: {json.dumps(interactions[-20:], ensure_ascii=False)[:6000]}\n"
            "Respond in 2-4 sentences and reference evidence from the interactions."
        )
        try:
            return self.llm.complete_required(
                prompt,
                system_prompt="You are McKAInsey ReportAgent. Stay factual and evidence-grounded.",
            )
        except Exception:  # noqa: BLE001
            return "The available interactions indicate this question can be answered from observed cohort arguments and sentiment shifts."

    def _build_demographic_breakdown(self, agents: list[dict[str, Any]]) -> list[dict[str, Any]]:
        grouped: dict[str, dict[str, int]] = defaultdict(lambda: {"supporter": 0, "neutral": 0, "dissenter": 0})
        for agent in agents:
            segment = str(agent.get("persona", {}).get("planning_area", "Unknown"))
            score = float(agent.get("opinion_post", 5.0) or 5.0)
            if score >= 7:
                grouped[segment]["supporter"] += 1
            elif score >= 5:
                grouped[segment]["neutral"] += 1
            else:
                grouped[segment]["dissenter"] += 1
        rows = [
            {
                "segment": segment,
                "supporter": values["supporter"],
                "neutral": values["neutral"],
                "dissenter": values["dissenter"],
            }
            for segment, values in grouped.items()
        ]
        rows.sort(key=lambda row: row["dissenter"], reverse=True)
        return rows

    def _build_v2_recommendations(self, demographic_breakdown: list[dict[str, Any]], dissenting_views: list[str]) -> list[str]:
        recommendations: list[str] = []
        if demographic_breakdown:
            top = demographic_breakdown[0]
            recommendations.append(
                f"Prioritize communication and safeguards for {top.get('segment', 'top dissent segment')} to reduce concentrated dissent."
            )
        if dissenting_views:
            recommendations.append("Address recurring affordability concerns directly with concrete implementation details.")
        if not recommendations:
            recommendations.append("Maintain transparent rollout updates and monitor stance movement each round.")
        return recommendations[:5]

    def _build_v2_executive_summary(
        self,
        *,
        simulation_id: str,
        initial_metric: float,
        final_metric: float,
        round_count: int,
        supporting_views: list[str],
        dissenting_views: list[str],
    ) -> str:
        prompt = (
            f"Simulation {simulation_id}. Approval moved from {initial_metric} to {final_metric} "
            f"over {round_count} rounds.\n"
            f"Supporting themes: {supporting_views[:3]}\n"
            f"Dissenting themes: {dissenting_views[:3]}\n"
            "Write a concise executive summary in 3-4 sentences."
        )
        try:
            return self.llm.complete_required(prompt, system_prompt="You are McKAInsey ReportAgent.")
        except Exception:  # noqa: BLE001
            direction = "declined" if final_metric < initial_metric else "improved"
            return (
                f"Across {round_count} rounds, overall approval {direction} from {initial_metric} to {final_metric}. "
                "Observed interactions show concentrated disagreement around affordability and rollout fairness."
            )

    # ── New V2 helper methods ──

    def _session_analysis_questions(self, simulation_id: str) -> list[dict[str, Any]]:
        try:
            conn = sqlite3.connect(self.settings.simulation_db_path)
            row = conn.execute(
                "SELECT analysis_questions FROM session_configs WHERE session_id = ?",
                (simulation_id,),
            ).fetchone()
            conn.close()
        except Exception:  # noqa: BLE001
            return []
        if not row or row[0] is None:
            return []
        raw = row[0]
        if isinstance(raw, str):
            try:
                raw = json.loads(raw)
            except Exception:  # noqa: BLE001
                return []
        if not isinstance(raw, list):
            return []
        return [item for item in raw if isinstance(item, dict)]

    def _resolve_analysis_questions(self, *, simulation_id: str, use_case: str | None) -> list[dict[str, Any]]:
        """Resolve analysis questions from session config first, then fall back to use-case YAML."""
        session_questions = self._session_analysis_questions(simulation_id)
        if session_questions:
            return session_questions

        if use_case:
            config_service = ConfigService(self.settings)
            try:
                questions = config_service.get_analysis_questions(use_case)
                if questions:
                    return questions
                sections = config_service.get_report_sections(use_case)
                normalized_from_sections = [
                    {
                        "question": str(item.get("prompt") or item.get("title") or "").strip(),
                        "type": "open-ended",
                        "metric_name": str(item.get("title", "")).strip().lower().replace(" ", "_") or "section",
                        "report_title": str(item.get("title") or item.get("prompt") or "Section").strip(),
                        "tooltip": "",
                    }
                    for item in sections
                    if isinstance(item, dict) and str(item.get("prompt") or item.get("title") or "").strip()
                ]
                if normalized_from_sections:
                    return normalized_from_sections
            except Exception:  # noqa: BLE001
                pass
        return []

    def _resolve_insight_blocks(self, use_case: str | None) -> list[dict[str, Any]]:
        """Resolve insight block configs from the use-case YAML."""
        if use_case:
            config_service = ConfigService(self.settings)
            try:
                return config_service.get_insight_blocks(use_case)
            except Exception:  # noqa: BLE001
                pass
        return []

    def _resolve_preset_sections(self, use_case: str | None) -> list[dict[str, Any]]:
        """Resolve preset section configs from the use-case YAML."""
        if use_case:
            config_service = ConfigService(self.settings)
            try:
                return config_service.get_preset_sections(use_case)
            except Exception:  # noqa: BLE001
                pass
        return []

    def _agents_from_checkpoint(self, checkpoint_records: list[dict[str, Any]]) -> list[dict[str, Any]] | None:
        """Convert checkpoint records into agent-like dicts for metric computation."""
        if not checkpoint_records:
            return None
        rows: list[dict[str, Any]] = []
        for record in checkpoint_records:
            if not isinstance(record, dict):
                continue
            row: dict[str, Any] = {}
            metric_answers = record.get("metric_answers")
            if isinstance(metric_answers, dict):
                for metric_name, value in metric_answers.items():
                    clean_name = str(metric_name or "").strip()
                    if not clean_name:
                        continue
                    row[f"checkpoint_{clean_name}"] = value
            if row:
                rows.append(row)
            else:
                rows.append(record)
        return rows

    def _compute_metric_value(self, question: dict[str, Any], agents: list[dict[str, Any]]) -> float:
        """Compute a single metric value for one analysis question across agents."""
        name = question.get("metric_name", "")
        field = f"checkpoint_{name}"
        q_type = question.get("type", "scale")
        total = max(len(agents), 1)

        if q_type == "scale":
            scores = [float(a.get(field, a.get("opinion_post", 5.0)) or 5.0) for a in agents]
            if "threshold" in question:
                threshold = float(question.get("threshold", 7))
                pct = sum(1 for s in scores if s >= threshold) / total * 100
                return round(pct, 1)
            return round(sum(scores) / len(scores) if scores else 0.0, 1)
        elif q_type == "yes-no":
            yes_count = sum(1 for a in agents if str(a.get(field, "")).strip().lower() in {"yes", "y"})
            return round(yes_count / total * 100, 1)
        return 0.0

    def _build_v2_executive_summary_from_metrics(
        self,
        *,
        simulation_id: str,
        metric_deltas: list[dict[str, Any]],
        round_count: int,
        agent_count: int,
    ) -> str:
        """Build executive summary using metric deltas instead of raw scores."""
        if not metric_deltas:
            return f"Simulation {simulation_id} completed with {agent_count} agents over {round_count} rounds."

        metrics_summary = "; ".join(
            f"{d['metric_label']}: {d['initial_display']} -> {d['final_display']} ({'+' if d['delta'] > 0 else ''}{d['delta']}{d['metric_unit']})"
            for d in metric_deltas
        )
        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}
        knowledge_context = "\n".join(self._knowledge_context_lines(knowledge))
        prompt = (
            f"Simulation {simulation_id}. {agent_count} agents over {round_count} rounds.\n"
            f"Original document context:\n{knowledge_context or '- none'}\n"
            f"Key metrics: {metrics_summary}\n"
            "Write a concise executive summary in 3-4 sentences highlighting the most important findings."
        )
        try:
            return self.llm.complete_required(prompt, system_prompt="You are McKAInsey ReportAgent.")
        except Exception:  # noqa: BLE001
            return f"Across {round_count} rounds with {agent_count} agents: {metrics_summary}."

    def _recommend(
        self,
        simulation_id: str,
        top_dissenting: list[dict[str, Any]],
        income_metrics: list[dict[str, Any]],
        arguments_for: list[dict[str, Any]],
        arguments_against: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        if not top_dissenting:
            return [
                {
                    "title": "Maintain broad-based communication cadence",
                    "rationale": "No major friction clusters detected in planning-area analysis.",
                    "target_demographic": "All cohorts",
                    "expected_impact": "Medium",
                    "execution_plan": [
                        "Keep monthly policy updates with simple impact examples.",
                        "Run sentiment pulse checks by demographic cohorts.",
                    ],
                    "confidence": 0.62,
                }
            ]

        prompt = (
            "Generate 5 concrete policy communication/mitigation recommendations in JSON. "
            "Use ONLY this schema: "
            "[{\"title\": str, \"rationale\": str, \"target_demographic\": str, "
            "\"expected_impact\": str, \"execution_plan\": [str, str, str], \"confidence\": number}]\n"
            f"simulation_id={simulation_id}\n"
            f"top_dissenting={top_dissenting[:6]}\n"
            f"income_metrics={sorted(income_metrics, key=lambda x: x['approval_post'])[:6]}\n"
            f"arguments_for={arguments_for[:6]}\n"
            f"arguments_against={arguments_against[:6]}\n"
            "Rules: recommendations must be specific, non-generic, and tied to at least one planning area or cohort. "
            "confidence must be between 0 and 1."
        )

        raw = self.llm.complete_required(
            prompt=prompt,
            system_prompt="You are McKAInsey ReportAgent. Return valid JSON only.",
        )
        parsed = self._parse_recommendations(raw)
        if parsed:
            return parsed
        raise RuntimeError("Report recommendation generation failed because the model did not return valid JSON.")

    def _parse_recommendations(self, raw: str) -> list[dict[str, Any]]:
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.strip("`")
            if cleaned.lower().startswith("json"):
                cleaned = cleaned[4:].strip()

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            return []

        if not isinstance(data, list):
            return []

        out: list[dict[str, Any]] = []
        for item in data:
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            rationale = str(item.get("rationale", "")).strip()
            target = str(item.get("target_demographic", "")).strip()
            impact = str(item.get("expected_impact", "")).strip()
            plan = item.get("execution_plan", [])
            try:
                conf = float(item.get("confidence", 0.5))
            except (TypeError, ValueError):
                conf = 0.5

            if not title or not rationale or not target:
                continue

            plan_list = [str(x).strip() for x in plan if str(x).strip()]
            if len(plan_list) < 2:
                plan_list = [
                    "Run targeted messaging sessions with affected households.",
                    "Track sentiment changes weekly and refine intervention messaging.",
                ]

            out.append(
                {
                    "title": title,
                    "rationale": rationale,
                    "target_demographic": target,
                    "expected_impact": impact or "Medium",
                    "execution_plan": plan_list[:4],
                    "confidence": max(0.0, min(1.0, round(conf, 2))),
                }
            )

        return out[:6]

    def _algorithmic_recommendations(
        self,
        top_dissenting: list[dict[str, Any]],
        income_metrics: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        recommendations: list[dict[str, Any]] = []
        low_income = sorted(income_metrics, key=lambda x: x["approval_post"])[:2]

        for item in top_dissenting[:4]:
            area = item["planning_area"]
            friction = float(item.get("friction_index", 0.0))
            target_income = low_income[0]["income_bracket"] if low_income else "Lower-income households"
            confidence = 0.55 + min(0.35, friction)
            recommendations.append(
                {
                    "title": f"Targeted affordability mitigation for {area}",
                    "rationale": (
                        f"{area} shows elevated friction ({friction:.2f}) with below-target post approval "
                        f"({item.get('approval_post', 0):.2f})."
                    ),
                    "target_demographic": f"{area} residents, especially {target_income}",
                    "expected_impact": "High" if friction >= 0.3 else "Medium",
                    "execution_plan": [
                        f"Deploy area-specific budget explainers in {area} community channels.",
                        "Add concrete household cashflow examples for affected segments.",
                        "Collect 2-week feedback pulse and adjust subsidy messaging.",
                    ],
                    "confidence": round(min(0.95, confidence), 2),
                }
            )

        if not recommendations:
            recommendations.append(
                {
                    "title": "Cross-cohort message calibration",
                    "rationale": "No sharply concentrated friction cluster was detected.",
                    "target_demographic": "Multi-cohort",
                    "expected_impact": "Medium",
                    "execution_plan": [
                        "Segment messages by age and income before public rollout.",
                        "Prioritize FAQs around transport and cost-of-living concerns.",
                    ],
                    "confidence": 0.6,
                }
            )

        return recommendations[:6]

    def _build_structured_report_prompt(
        self,
        *,
        simulation_id: str,
        use_case: str | None,
        knowledge: dict[str, Any],
        population: dict[str, Any],
        agents: list[dict[str, Any]],
        interactions: list[dict[str, Any]],
        baseline: list[dict[str, Any]],
        final: list[dict[str, Any]],
        events: list[dict[str, Any]],
    ) -> str:
        config_lines: list[str] = []
        if use_case:
            config_service = ConfigService(self.settings)
            try:
                use_case_payload = config_service.get_use_case(use_case)
            except Exception:  # noqa: BLE001
                use_case_payload = {}
            guiding_prompt = str(use_case_payload.get("guiding_prompt") or "").strip()
            if guiding_prompt:
                config_lines.append("Use-case guiding prompt:")
                config_lines.append(guiding_prompt)
            report_sections = [
                item
                for item in use_case_payload.get("report_sections", [])
                if isinstance(item, dict)
            ]
            if report_sections:
                config_lines.append("Report sections from config:")
                for index, section in enumerate(report_sections, start=1):
                    title = str(section.get("title") or "").strip()
                    prompt = str(section.get("prompt") or "").strip()
                    if title or prompt:
                        config_lines.append(f"{index}. {title}: {prompt}".strip())
        influential_posts = [
            {
                "agent_id": row.get("actor_agent_id"),
                "content": row.get("content"),
                "delta": row.get("delta"),
            }
            for row in interactions
            if row.get("action_type") == "create_post"
        ][:12]
        checkpoints = {
            "baseline": baseline[:50],
            "final": final[:50],
        }
        prompt_lines = [
            "Generate a fixed-format policy simulation report in JSON.",
            "Return an object with exactly these top-level keys:",
            "{\"generated_at\": str, \"executive_summary\": str, "
            "\"insight_cards\": [{\"title\": str, \"summary\": str, \"severity\": \"high|medium|low\"}], "
            "\"support_themes\": [{\"theme\": str, \"summary\": str, \"evidence\": [str]}], "
            "\"dissent_themes\": [{\"theme\": str, \"summary\": str, \"evidence\": [str]}], "
            "\"demographic_breakdown\": [{\"segment\": str, \"approval_rate\": number, \"dissent_rate\": number, \"sample_size\": number}], "
            "\"influential_content\": [{\"content_type\": str, \"author_agent_id\": str, \"summary\": str, \"engagement_score\": number}], "
            "\"recommendations\": [{\"title\": str, \"rationale\": str, \"priority\": \"high|medium|low\"}], "
            "\"risks\": [{\"title\": str, \"summary\": str, \"severity\": \"high|medium|low\"}]}",
            "",
        ]
        if config_lines:
            prompt_lines.extend(config_lines)
            prompt_lines.append("")
        prompt_lines.extend(
            [
                f"Simulation ID: {simulation_id}",
                f"Knowledge summary: {knowledge.get('summary', '')}",
                f"Population artifact: {json.dumps(population, ensure_ascii=False)[:6000]}",
                f"Checkpoint records: {json.dumps(checkpoints, ensure_ascii=False)[:12000]}",
                f"Influential posts: {json.dumps(influential_posts, ensure_ascii=False)[:6000]}",
                f"Recent simulation events: {json.dumps(events[-80:], ensure_ascii=False)[:12000]}",
                f"Agent records: {json.dumps(agents[:80], ensure_ascii=False)[:12000]}",
            ]
        )
        return "\n".join(prompt_lines)

    def _normalize_structured_report_payload(self, simulation_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        generated_at = str(payload.get("generated_at") or datetime.now(UTC).isoformat())
        normalized = {
            "session_id": simulation_id,
            "status": "completed",
            "generated_at": generated_at,
            "executive_summary": str(payload.get("executive_summary", "")).strip(),
            "insight_cards": _normalize_dict_list(payload.get("insight_cards"), required_keys=("title", "summary", "severity")),
            "support_themes": _normalize_dict_list(payload.get("support_themes"), required_keys=("theme", "summary", "evidence")),
            "dissent_themes": _normalize_dict_list(payload.get("dissent_themes"), required_keys=("theme", "summary", "evidence")),
            "demographic_breakdown": _normalize_dict_list(payload.get("demographic_breakdown"), required_keys=("segment", "approval_rate", "dissent_rate", "sample_size")),
            "influential_content": _normalize_dict_list(payload.get("influential_content"), required_keys=("content_type", "author_agent_id", "summary", "engagement_score")),
            "recommendations": _normalize_dict_list(payload.get("recommendations"), required_keys=("title", "rationale", "priority")),
            "risks": _normalize_dict_list(payload.get("risks"), required_keys=("title", "summary", "severity")),
        }
        return normalized

    def _enrich_structured_report_payload(
        self,
        simulation_id: str,
        payload: dict[str, Any],
        *,
        use_case: str | None,
        agents: list[dict[str, Any]],
        interactions: list[dict[str, Any]],
        baseline: list[dict[str, Any]],
        final: list[dict[str, Any]],
        events: list[dict[str, Any]],
        knowledge: dict[str, Any],
        population: dict[str, Any],
    ) -> dict[str, Any]:
        enriched = dict(payload)
        pre_scores = [float(agent.get("opinion_pre", 5.0) or 5.0) for agent in agents]
        post_scores = [float(agent.get("opinion_post", 5.0) or 5.0) for agent in agents]
        approval_pre = _approval(pre_scores)
        approval_post = _approval(post_scores)

        supportive_rows = self._rank_interactions(interactions, positive=True)
        dissent_rows = self._rank_interactions(interactions, positive=False)
        demographic_breakdown = enriched.get("demographic_breakdown") or self._build_demographic_breakdown(agents)

        if not enriched["executive_summary"]:
            enriched["executive_summary"] = self._build_structured_executive_summary(
                simulation_id=simulation_id,
                use_case=use_case,
                demographic_breakdown=demographic_breakdown,
                supportive_rows=supportive_rows,
                dissent_rows=dissent_rows,
                approval_pre=approval_pre,
                approval_post=approval_post,
            )

        if not enriched["insight_cards"]:
            top_segment = str(demographic_breakdown[0].get("segment", "top cohort")) if demographic_breakdown else "top cohort"
            card_summary = (
                f"{top_segment} carried the strongest signal in the simulation, "
                f"with approval moving from {approval_pre:.2f} to {approval_post:.2f} across the run."
            )
            enriched["insight_cards"] = [
                {
                    "title": f"{top_segment} drove the clearest shift",
                    "summary": card_summary,
                    "severity": "high" if abs(approval_post - approval_pre) >= 0.15 else "medium",
                }
            ]
            if supportive_rows:
                first_support = supportive_rows[0]
                enriched["insight_cards"].append(
                    {
                        "title": "Most persuasive support argument",
                        "summary": str(first_support["content"])[:240],
                        "severity": "medium",
                    }
                )
            if dissent_rows:
                first_dissent = dissent_rows[0]
                enriched["insight_cards"].append(
                    {
                        "title": "Main dissent pressure point",
                        "summary": str(first_dissent["content"])[:240],
                        "severity": "medium",
                    }
                )

        if not enriched["support_themes"]:
            enriched["support_themes"] = self._build_theme_items(
                supportive_rows,
                theme_label="support",
                fallback_summary="Support centered on concrete benefits and targeted help.",
            )

        if not enriched["dissent_themes"]:
            enriched["dissent_themes"] = self._build_theme_items(
                dissent_rows,
                theme_label="dissent",
                fallback_summary="Dissent clustered around affordability, fairness, or implementation risk.",
            )

        if not enriched["demographic_breakdown"]:
            enriched["demographic_breakdown"] = demographic_breakdown

        if not enriched["influential_content"]:
            enriched["influential_content"] = self._build_influential_content(interactions)

        if not enriched["recommendations"]:
            enriched["recommendations"] = self._build_structured_recommendations(
                simulation_id=simulation_id,
                demographic_breakdown=demographic_breakdown,
                dissent_rows=dissent_rows,
                supportive_rows=supportive_rows,
                knowledge=knowledge,
                population=population,
                use_case=use_case,
                baseline=baseline,
                final=final,
                events=events,
            )

        if not enriched["risks"]:
            enriched["risks"] = self._build_structured_risks(
                demographic_breakdown=demographic_breakdown,
                dissent_rows=dissent_rows,
                events=events,
            )

        return enriched

    def _rank_interactions(self, interactions: list[dict[str, Any]], *, positive: bool) -> list[dict[str, Any]]:
        rows: list[dict[str, Any]] = []
        for item in interactions:
            try:
                delta = float(item.get("delta", 0.0) or 0.0)
            except (TypeError, ValueError):
                delta = 0.0
            if positive and delta <= 0:
                continue
            if not positive and delta >= 0:
                continue
            content = str(item.get("content", "")).strip()
            if not content:
                continue
            rows.append(
                {
                    "content": content,
                    "agent_id": str(item.get("actor_agent_id", "")),
                    "round_no": int(item.get("round_no", 0) or 0),
                    "delta": delta,
                    "likes": float(item.get("likes", 0) or 0),
                    "dislikes": float(item.get("dislikes", 0) or 0),
                }
            )
        rows.sort(key=lambda row: (abs(row["delta"]), row["likes"] + row["dislikes"]), reverse=True)
        return rows[:6]

    def _build_theme_items(
        self,
        rows: list[dict[str, Any]],
        *,
        theme_label: str,
        fallback_summary: str,
    ) -> list[dict[str, Any]]:
        if not rows:
            return [
                {
                    "theme": theme_label,
                    "summary": fallback_summary,
                    "evidence": [],
                }
            ]
        items: list[dict[str, Any]] = []
        for row in rows[:3]:
            summary = f"{row['content'][:180]}"
            items.append(
                {
                    "theme": theme_label,
                    "summary": summary,
                    "evidence": [row["content"]],
                }
            )
        return items

    def _build_influential_content(self, interactions: list[dict[str, Any]]) -> list[dict[str, Any]]:
        rows = []
        for item in interactions:
            content = str(item.get("content", "")).strip()
            if not content:
                continue
            try:
                delta = abs(float(item.get("delta", 0.0) or 0.0))
            except (TypeError, ValueError):
                delta = 0.0
            try:
                likes = float(item.get("likes", 0) or 0)
            except (TypeError, ValueError):
                likes = 0.0
            try:
                dislikes = float(item.get("dislikes", 0) or 0)
            except (TypeError, ValueError):
                dislikes = 0.0
            engagement_score = round(delta * 10 + likes + dislikes, 2)
            rows.append(
                {
                    "content_type": str(item.get("action_type") or item.get("type") or "post"),
                    "author_agent_id": str(item.get("actor_agent_id", "")),
                    "summary": content[:240],
                    "engagement_score": engagement_score,
                }
            )
        rows.sort(key=lambda row: row["engagement_score"], reverse=True)
        return rows[:6]

    def _build_structured_recommendations(
        self,
        *,
        simulation_id: str,
        demographic_breakdown: list[dict[str, Any]],
        dissent_rows: list[dict[str, Any]],
        supportive_rows: list[dict[str, Any]],
        knowledge: dict[str, Any],
        population: dict[str, Any],
        use_case: str | None,
        baseline: list[dict[str, Any]],
        final: list[dict[str, Any]],
        events: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        top_segment = str(demographic_breakdown[0].get("segment", "All cohorts")) if demographic_breakdown else "All cohorts"
        top_dissent = dissent_rows[0]["content"] if dissent_rows else "review implementation gaps"
        support_context = supportive_rows[0]["content"] if supportive_rows else str(knowledge.get("summary", "")).strip()
        base_label = use_case or str(population.get("use_case") or "simulation")
        return [
            {
                "title": f"Address the main friction in {top_segment}",
                "rationale": f"Dissent in {top_segment} is the clearest signal to act on first.",
                "priority": "high",
            },
            {
                "title": f"Turn the strongest support into a clearer message for {base_label}",
                "rationale": support_context[:240] or "Support needs to be translated into a more concrete narrative.",
                "priority": "medium",
            },
            {
                "title": "Use round-by-round evidence to close credibility gaps",
                "rationale": top_dissent[:240] if top_dissent else "Agents responded to concrete examples more than abstract assurances.",
                "priority": "medium",
            },
        ]

    def _build_structured_risks(
        self,
        *,
        demographic_breakdown: list[dict[str, Any]],
        dissent_rows: list[dict[str, Any]],
        events: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        risks: list[dict[str, Any]] = []
        if demographic_breakdown:
            top = demographic_breakdown[0]
            risks.append(
                {
                    "title": f"Concentrated dissent in {top.get('segment', 'a key cohort')}",
                    "summary": (
                        f"{top.get('segment', 'A cohort')} has {top.get('dissent_rate', 0)} dissent rate "
                        f"across {top.get('sample_size', 0)} agents."
                    ),
                    "severity": "high" if float(top.get("dissent_rate", 0) or 0) >= 0.3 else "medium",
                }
            )
        if dissent_rows:
            risks.append(
                {
                    "title": "Recurring objection pattern",
                    "summary": dissent_rows[0]["content"][:240],
                    "severity": "medium",
                }
            )
        if events:
            risks.append(
                {
                    "title": "Conversation may be dominated by the most active agents",
                    "summary": "Event logs show the report is driven by a small set of highly visible posts.",
                    "severity": "low",
                }
            )
        return risks[:4]

    def _build_structured_executive_summary(
        self,
        *,
        simulation_id: str,
        use_case: str | None,
        demographic_breakdown: list[dict[str, Any]],
        supportive_rows: list[dict[str, Any]],
        dissent_rows: list[dict[str, Any]],
        approval_pre: float,
        approval_post: float,
    ) -> str:
        top_segment = str(demographic_breakdown[0].get("segment", "the main cohort")) if demographic_breakdown else "the main cohort"
        support_excerpt = supportive_rows[0]["content"][:140] if supportive_rows else "support stayed concentrated in a few concrete arguments"
        dissent_excerpt = dissent_rows[0]["content"][:140] if dissent_rows else "dissent stayed centered on implementation risk"
        direction = "improved" if approval_post >= approval_pre else "softened"
        use_case_label = f"for {use_case}" if use_case else "for the simulation"
        return (
            f"Across {use_case_label}, approval {direction} from {approval_pre:.2f} to {approval_post:.2f}. "
            f"{top_segment} was the clearest cohort signal in the run, with support anchored by '{support_excerpt}' "
            f"and dissent concentrated around '{dissent_excerpt}'. "
            "The report sections point to a need for sharper mitigation and clearer rollout messaging."
        )


def _approval(scores: list[float]) -> float:
    if not scores:
        return 0.0
    return len([s for s in scores if s >= 7]) / len(scores)


def _mean(scores: list[float]) -> float:
    if not scores:
        return 0.0
    return sum(scores) / len(scores)


def _parse_json_object(raw: str) -> Any:
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:].strip()
    return json.loads(cleaned)


def _normalize_dict_list(value: Any, *, required_keys: tuple[str, ...]) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    normalized: list[dict[str, Any]] = []
    for item in value:
        if not isinstance(item, dict):
            continue
        normalized.append({key: item.get(key) for key in required_keys})
    return normalized
