from __future__ import annotations

from datetime import UTC, datetime
import io
import json
import re
from typing import Any

from docx import Document

from miroworld.config import Settings
from miroworld.services.config_service import ConfigService
from miroworld.services.llm_client import GeminiChatClient
from miroworld.services.memory_service import MemoryService
from miroworld.services.metrics_service import MetricsService, build_influence_graph, compute_opinion_flow, compute_polarization
from miroworld.services.storage import SimulationStore


def _clean_report_text(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"[*_`#]+", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _clean_bullet_text(value: Any) -> str:
    text = _clean_report_text(value)
    text = re.sub(r"^(?:[-*•]+|\d+[.)])\s*", "", text)
    return text.strip()


def _format_metric_value(value: float, unit: str) -> str:
    if unit == "%":
        return f"{value:.1f}%"
    if unit == "/10":
        return f"{value:.1f}/10"
    return f"{value:.1f}"


def _extract_numeric_value(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)

    text = str(value or "").strip()
    if not text:
        return None

    ratio_match = re.search(r"(-?\d+(?:\.\d+)?)\s*/\s*(-?\d+(?:\.\d+)?)", text)
    if ratio_match:
        numerator = float(ratio_match.group(1))
        denominator = float(ratio_match.group(2))
        if denominator != 0:
            return numerator if denominator == 10 else (numerator / denominator) * 10

    number_match = re.search(r"-?\d+(?:\.\d+)?", text)
    if number_match:
        return float(number_match.group(0))

    return None


def _parse_yes_no(value: Any) -> bool | None:
    if isinstance(value, bool):
        return value

    text = str(value or "").strip().lower()
    if not text:
        return None

    if re.match(r"^(yes|y|true|1)\b", text):
        return True
    if re.match(r"^(no|n|false|0)\b", text):
        return False
    return None


def _parse_json_value(raw: str) -> Any:
    cleaned = str(raw or "").strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    return json.loads(cleaned)


def _contains_first_person(text: str) -> bool:
    lowered = f" {str(text or '').lower()} "
    return any(token in lowered for token in (" i ", " i'm ", " i've ", " i'd ", " we ", " we're ", " we've ", " our ", " my "))


def _render_prompt_template(template: str, **values: Any) -> str:
    rendered = str(template or "")
    for key, value in values.items():
        rendered = rendered.replace(f"{{{key}}}", str(value))
    return rendered.replace("{{", "{").replace("}}", "}")


def _normalize_bullet_payload(
    payload: Any,
    *,
    min_bullets: int,
    max_bullets: int,
    max_words_per_bullet: int,
) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    if isinstance(payload, dict):
        raw_bullets = payload.get("bullets")
    else:
        raw_bullets = payload

    if not isinstance(raw_bullets, list):
        return [], ["Response must be a JSON object with a bullets array."]

    bullets = [_clean_bullet_text(item) for item in raw_bullets if _clean_bullet_text(item)]
    if len(bullets) < min_bullets:
        errors.append(f"Return at least {min_bullets} bullets.")
    if len(bullets) > max_bullets:
        errors.append(f"Return no more than {max_bullets} bullets.")

    for bullet in bullets:
        if _contains_first_person(bullet):
            errors.append("Bullets must stay in third-person analytical voice.")
        if len(bullet.split()) > max_words_per_bullet:
            errors.append(f"Each bullet must stay within {max_words_per_bullet} words.")

    return bullets[:max_bullets], errors


def _bullets_to_text(bullets: list[str]) -> str:
    return " ".join(_clean_bullet_text(item) for item in bullets if _clean_bullet_text(item)).strip()


class ReportService:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.store = SimulationStore(settings.simulation_db_path)
        self.llm = GeminiChatClient(settings)
        self.memory = MemoryService(settings)
        self.config = ConfigService(settings)

    def report_chat_payload(self, simulation_id: str, message: str) -> dict[str, Any]:
        report = self.build_v2_report(simulation_id)
        memory_context = self.memory.search_simulation_context(simulation_id, message, limit=8)
        knowledge = self.store.get_knowledge_artifact(simulation_id) or {}
        memory_excerpt = self._format_memory_excerpt(memory_context["episodes"])
        checkpoint_excerpt = self.memory.format_checkpoint_records(memory_context["checkpoint_records"], limit=6)
        knowledge_excerpt = "\n".join(self._knowledge_context_lines(knowledge))
        prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "report_chat",
            "user_template_with_memory",
        ).format(
            report_json=json.dumps(report, ensure_ascii=False),
            knowledge_excerpt=knowledge_excerpt or "- none",
            memory_excerpt=memory_excerpt or "- none",
            checkpoint_excerpt=checkpoint_excerpt or "- none",
            message=message,
        )
        response = self.llm.complete_required(
            prompt,
            system_prompt=self.config.get_system_prompt_value(
                "report_agent",
                "prompts",
                "report_chat",
                "system_prompt",
            ),
        )
        return {
            "session_id": simulation_id,
            "simulation_id": simulation_id,
            "response": response,
            "model_provider": self.llm.provider,
            "model_name": self.llm.model_name,
            "gemini_model": self.llm.model_name,
            "zep_context_used": bool(memory_context.get("zep_context_used", False)),
            "graphiti_context_used": False,
            "memory_backend": memory_context.get("memory_backend", self.memory.memory_backend),
        }

    def build_v2_report(self, simulation_id: str, use_case: str | None = None) -> dict[str, Any]:
        agents = self.store.get_agents(simulation_id)
        interactions = self.store.get_interactions(simulation_id)
        if not agents:
            raise ValueError(f"Simulation not found: {simulation_id}")

        resolved_use_case = use_case or self._session_use_case(simulation_id)
        metrics_service = MetricsService(self.config)
        analysis_questions = self._resolve_analysis_questions(simulation_id=simulation_id, use_case=resolved_use_case)
        insight_block_configs = self._resolve_insight_blocks(resolved_use_case)
        preset_section_configs = self._resolve_preset_sections(resolved_use_case)

        simulation = self.store.get_simulation(simulation_id) or {}
        stored_round_count = int(simulation.get("rounds", 0) or 0)
        interaction_round_count = max((int(item.get("round_no", 0) or 0) for item in interactions), default=0)
        round_count = stored_round_count or interaction_round_count
        baseline_records = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="baseline")
        final_records = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="final")
        metric_deltas: list[dict[str, Any]] = []
        for question in analysis_questions:
            if question.get("type") == "open-ended":
                continue
            metric_name = str(question.get("metric_name", "")).strip()
            if not metric_name:
                continue

            baseline_agents = self._agents_from_checkpoint(baseline_records) or agents
            final_agents = self._agents_from_checkpoint(final_records) or agents
            initial_value = self._compute_metric_value(question, baseline_agents)
            final_value = self._compute_metric_value(question, final_agents)
            metric_unit = self._metric_unit(question)
            initial_display = _format_metric_value(initial_value, metric_unit)
            final_display = _format_metric_value(final_value, metric_unit)
            metric_deltas.append(
                {
                    "metric_name": metric_name,
                    "metric_label": question.get("metric_label", metric_name),
                    "metric_unit": metric_unit,
                    "initial_value": initial_value,
                    "final_value": final_value,
                    "delta": round(final_value - initial_value, 2),
                    "direction": "up" if final_value > initial_value else ("down" if final_value < initial_value else "flat"),
                    "report_title": question.get("report_title", metric_name),
                    "initial_display": initial_display,
                    "final_display": final_display,
                    "delta_display": f"{initial_display} -> {final_display}",
                    "type": question.get("type", "scale"),
                }
            )

        sections: list[dict[str, Any]] = []
        for question in analysis_questions:
            question_text = str(question.get("question", "")).strip()
            question_type = str(question.get("type", "scale")).strip()
            delta_entry = None
            if question_type != "open-ended":
                delta_entry = next((item for item in metric_deltas if item["metric_name"] == question.get("metric_name")), None)
            raw_bullets, context_pack = self._answer_guiding_question(
                simulation_id,
                question_text,
                agent_count=len(agents),
                question_type=question_type,
                metric=delta_entry,
                use_case=resolved_use_case,
                section_type="analysis_section",
                section_title=str(question.get("report_title") or question_text).strip(),
            )
            pack_evidence = list(context_pack.get("top_discourse_evidence") or [])
            section_evidence = self._select_section_evidence(
                question_text,
                _bullets_to_text(raw_bullets),
                pack_evidence,
                limit=4,
            )
            if not section_evidence:
                section_evidence = pack_evidence[:4]
            display_bullets = self._replace_references_in_bullets(raw_bullets, section_evidence)
            section: dict[str, Any] = {
                "question": question_text,
                "report_title": question.get("report_title", question_text[:60]),
                "type": question_type,
                "bullets": display_bullets,
                "evidence": section_evidence,
            }
            if delta_entry:
                section["metric"] = delta_entry
            sections.append(section)

        insight_blocks: list[dict[str, Any]] = []
        for block_cfg in insight_block_configs:
            block_type = str(block_cfg.get("type", "")).strip()
            if not block_type:
                continue
            result = metrics_service.compute_insight_block(
                block_type=block_type,
                agents=agents,
                interactions=interactions,
                analysis_questions=analysis_questions,
                metric_ref=block_cfg.get("metric_ref"),
                count=block_cfg.get("count", 5),
            )
            insight_blocks.append(
                {
                    "type": block_type,
                    "title": block_cfg.get("title", block_type),
                    "description": block_cfg.get("description", ""),
                    "data": result,
                }
            )

        preset_sections: list[dict[str, Any]] = []
        for preset in preset_section_configs:
            title = str(preset.get("title", "")).strip()
            prompt = str(preset.get("prompt", "")).strip()
            raw_bullets, context_pack = self._answer_guiding_question(
                simulation_id,
                prompt,
                agent_count=len(agents),
                question_type="preset",
                metric=None,
                use_case=resolved_use_case,
                section_type="preset_section",
                section_title=title or "Preset Section",
            )
            pack_evidence = list(context_pack.get("top_discourse_evidence") or [])
            section_evidence = self._select_section_evidence(
                title or prompt,
                _bullets_to_text(raw_bullets),
                pack_evidence,
                limit=4,
            )
            if not section_evidence:
                section_evidence = pack_evidence[:4]
            preset_sections.append(
                {
                    "title": title,
                    "bullets": self._replace_references_in_bullets(raw_bullets, section_evidence),
                }
            )

        executive_summary = self._build_v2_executive_summary_from_metrics(
            simulation_id=simulation_id,
            metric_deltas=metric_deltas,
            round_count=round_count,
            agent_count=len(agents),
            use_case=resolved_use_case,
        )

        return {
            "session_id": simulation_id,
            "status": "completed",
            "generated_at": datetime.now(UTC).isoformat(),
            "executive_summary": _clean_report_text(executive_summary),
            "metric_deltas": metric_deltas,
            "quick_stats": {
                "agent_count": len(agents),
                "round_count": round_count,
                "model": self.llm.model_name,
                "provider": self.llm.provider,
            },
            "sections": sections,
            "insight_blocks": insight_blocks,
            "preset_sections": preset_sections,
            "zep_context_used": self.memory.zep.enabled,
            "graphiti_context_used": False,
            "memory_backend": self.memory.memory_backend,
            "context_backend": "zep" if self.memory.zep.enabled else "sqlite",
            "context_pack_version": 1,
            "error": None,
        }

    def export_v2_report_docx(self, simulation_id: str, report: dict[str, Any] | None = None, use_case: str | None = None) -> bytes:
        payload = report or self.build_v2_report(simulation_id, use_case=use_case)
        document = Document()
        document.add_heading("MiroWorld Analysis Report", level=0)
        document.add_paragraph(f"Session: {payload.get('session_id', simulation_id)}")
        document.add_paragraph(f"Generated: {payload.get('generated_at', '')}")

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(str(payload.get("executive_summary", "")))

        quick_stats = payload.get("quick_stats", {})
        if isinstance(quick_stats, dict):
            document.add_heading("Quick Stats", level=1)
            document.add_paragraph(
                f"Agents: {quick_stats.get('agent_count', 0)} | Rounds: {quick_stats.get('round_count', 0)}"
            )
            document.add_paragraph(
                f"Model: {quick_stats.get('model', '')} ({quick_stats.get('provider', '')})"
            )

        metric_deltas = payload.get("metric_deltas", [])
        if isinstance(metric_deltas, list) and metric_deltas:
            document.add_heading("Metric Deltas", level=1)
            for metric in metric_deltas:
                if not isinstance(metric, dict):
                    continue
                label = str(metric.get("metric_label", metric.get("metric_name", "Metric")))
                initial_value = metric.get("initial_value", 0)
                final_value = metric.get("final_value", 0)
                delta = metric.get("delta", 0)
                unit = str(metric.get("metric_unit", ""))
                document.add_paragraph(
                    f"{label}: {initial_value}{unit} -> {final_value}{unit} ({delta:+}{unit})",
                    style="List Bullet",
                )

        document.add_heading("Analysis Findings", level=1)
        for section in payload.get("sections", []):
            if not isinstance(section, dict):
                continue
            section_title = str(section.get("report_title") or section.get("question") or "Section")
            document.add_heading(section_title, level=2)
            if section.get("question"):
                document.add_paragraph(f"Question: {section.get('question')}")
            for bullet in section.get("bullets", []):
                document.add_paragraph(str(bullet), style="List Bullet")
            evidence = section.get("evidence", [])
            if isinstance(evidence, list) and evidence:
                document.add_paragraph("Evidence:")
                for item in evidence:
                    if not isinstance(item, dict):
                        continue
                    agent_name = str(item.get("agent_name") or item.get("agent_id") or "Unknown agent")
                    quote = str(item.get("quote", "")).strip()
                    document.add_paragraph(f"{agent_name}: {quote}", style="List Bullet")

        preset_sections = payload.get("preset_sections", [])
        if isinstance(preset_sections, list) and preset_sections:
            document.add_heading("Preset Sections", level=1)
            for section in preset_sections:
                if not isinstance(section, dict):
                    continue
                document.add_heading(str(section.get("title", "Section")), level=2)
                for bullet in section.get("bullets", []):
                    document.add_paragraph(str(bullet), style="List Bullet")

        insight_blocks = payload.get("insight_blocks", [])
        if isinstance(insight_blocks, list) and insight_blocks:
            document.add_heading("Use-Case Insights", level=1)
            for block in insight_blocks:
                if not isinstance(block, dict):
                    continue
                document.add_heading(str(block.get("title", "Insight")), level=2)
                description = str(block.get("description", "")).strip()
                if description:
                    document.add_paragraph(description)

        try:
            agents_raw = self.store.get_agents(simulation_id)
            interactions = self.store.get_interactions(simulation_id)
            checkpoint_records = self.store.list_checkpoint_records(simulation_id, checkpoint_kind="post")
            if not checkpoint_records:
                checkpoint_records = self.store.list_checkpoint_records(simulation_id)

            metric_by_agent: dict[str, dict[str, Any]] = {}
            for record in checkpoint_records:
                aid = str(record.get("agent_id", "")).strip()
                if not aid:
                    continue
                answers = record.get("metric_answers") or {}
                if isinstance(answers, dict):
                    metric_by_agent.setdefault(aid, {}).update(answers)

            enriched_agents: list[dict[str, Any]] = []
            for row in agents_raw:
                agent = dict(row)
                agent["id"] = agent.get("id") or agent.get("agent_id")
                aid = str(agent.get("agent_id") or agent.get("id") or "")
                for metric_name, value in metric_by_agent.get(aid, {}).items():
                    agent[f"checkpoint_{metric_name}"] = value
                enriched_agents.append(agent)

            analysis_questions = self._resolve_analysis_questions(simulation_id=simulation_id, use_case=use_case or self._session_use_case(simulation_id))
            has_analytics = False
            for question in analysis_questions:
                if not isinstance(question, dict) or question.get("type") == "open-ended":
                    continue
                metric_name = str(question.get("metric_name", "")).strip()
                if not metric_name:
                    continue
                label = str(question.get("metric_label", metric_name)).strip()
                score_field = f"checkpoint_{metric_name}"
                polarization = compute_polarization(enriched_agents, score_field=score_field)
                opinion_flow = compute_opinion_flow(enriched_agents, score_field=score_field)
                if not has_analytics:
                    document.add_heading("Analytics Summary", level=1)
                    has_analytics = True
                document.add_heading(label, level=2)
                distribution = polarization.get("distribution", {})
                document.add_paragraph(
                    f"Polarization Index: {polarization.get('polarization_index', 0):.2f} ({polarization.get('severity', 'low')})",
                    style="List Bullet",
                )
                document.add_paragraph(
                    f"Distribution: {distribution.get('supporter_pct', 0):.0f}% supporters, "
                    f"{distribution.get('neutral_pct', 0):.0f}% neutral, "
                    f"{distribution.get('dissenter_pct', 0):.0f}% dissenters",
                    style="List Bullet",
                )
                initial = opinion_flow.get("initial", {})
                final = opinion_flow.get("final", {})
                document.add_paragraph(
                    f"Opinion Shift: Supporters {initial.get('supporter', 0)} -> {final.get('supporter', 0)}, "
                    f"Neutral {initial.get('neutral', 0)} -> {final.get('neutral', 0)}, "
                    f"Dissenters {initial.get('dissenter', 0)} -> {final.get('dissenter', 0)}",
                    style="List Bullet",
                )

            influence = build_influence_graph(interactions, enriched_agents)
            top_influencers = influence.get("top_influencers", [])[:5]
            if top_influencers:
                if not has_analytics:
                    document.add_heading("Analytics Summary", level=1)
                document.add_heading("Key Opinion Leaders", level=2)
                for influencer in top_influencers:
                    name_str = influencer.get("name", influencer.get("agent_id", ""))
                    stance = influencer.get("stance", "unknown")
                    score = influencer.get("influence_score", 0)
                    document.add_paragraph(
                        f"{name_str} ({stance}) - Influence: {score:.2f}",
                        style="List Bullet",
                    )
        except Exception:  # noqa: BLE001
            pass

        document.add_paragraph(
            f"Methodology: Simulated agents={quick_stats.get('agent_count', 0)}, "
            f"rounds={quick_stats.get('round_count', 0)}, "
            f"model={quick_stats.get('model', '')} ({quick_stats.get('provider', '')})."
        )

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _extract_evidence(self, interactions: list[dict[str, Any]], agents: list[dict[str, Any]]) -> list[dict[str, Any]]:
        agent_names: dict[str, str] = {}
        for agent in agents:
            agent_id = str(agent.get("agent_id") or agent.get("id") or "").strip()
            if not agent_id:
                continue
            persona = agent.get("persona") if isinstance(agent.get("persona"), dict) else {}
            display_name = _clean_report_text(
                agent.get("confirmed_name")
                or agent.get("name")
                or agent.get("agent_name")
                or agent.get("display_name")
                or agent.get("label")
                or persona.get("confirmed_name")
                or persona.get("name")
                or persona.get("agent_name")
                or persona.get("display_name")
                or persona.get("label")
                or agent_id
            )
            agent_names[agent_id] = display_name or agent_id

        discourse_actions = {
            "create_post",
            "post_created",
            "post",
            "comment",
            "comment_created",
            "reply",
        }
        evidence: list[dict[str, Any]] = []
        seen_quotes: set[str] = set()
        for row in interactions:
            action_type = str(row.get("action_type") or "").strip().lower()
            if action_type not in discourse_actions:
                continue

            quote = _clean_report_text(row.get("content") or row.get("body") or row.get("title"))
            if not quote or self._is_prompt_like_interaction(quote) or not self._is_opinion_like_interaction(quote):
                continue

            quote_key = quote.lower()
            if quote_key in seen_quotes:
                continue
            seen_quotes.add(quote_key)

            agent_id = str(row.get("actor_agent_id", "")).strip()
            agent_name = agent_names.get(agent_id, agent_id)
            source_type = "comment" if "comment" in action_type or "reply" in action_type else "post"
            evidence.append(
                {
                    "agent_id": agent_id,
                    "agent_name": agent_name,
                    "post_id": str(row.get("post_id") or row.get("id") or ""),
                    "action_type": action_type,
                    "source_type": source_type,
                    "source_label": f"{'Comment' if source_type == 'comment' else 'Post'} by {agent_name or 'Unknown agent'}",
                    "round_no": int(row.get("round_no", 0) or 0),
                    "quote": quote,
                }
            )
        return evidence

    def _tokenize_evidence_text(self, value: str) -> set[str]:
        stopwords = {
            "a", "an", "the", "and", "or", "to", "of", "for", "with", "in", "on", "at", "by",
            "is", "are", "was", "were", "be", "been", "being", "this", "that", "these", "those",
            "it", "its", "as", "from", "about", "into", "over", "under", "than", "then", "if",
            "but", "we", "our", "ours", "you", "your", "yours", "they", "their", "theirs",
            "agent", "policy", "persona", "perspective", "community", "prompt", "brief",
        }
        tokens = set(re.findall(r"[a-zA-Z][a-zA-Z0-9']+", str(value).lower()))
        return {token for token in tokens if len(token) >= 3 and token not in stopwords}

    def _is_prompt_like_interaction(self, text: str) -> bool:
        lowered = text.lower()
        if lowered.startswith("community prompt:"):
            return True
        if lowered.startswith("policy brief:"):
            return True
        if "community prompt:" in lowered and "policy brief:" in lowered:
            return True
        if (
            lowered.startswith("from your persona")
            or lowered.startswith("would you")
            or lowered.startswith("how strongly")
        ) and "?" in lowered:
            return True
        return False

    def _is_opinion_like_interaction(self, text: str) -> bool:
        lowered = text.lower().strip()
        if len(lowered.split()) < 8:
            return False
        opinion_markers = (
            "i ",
            "i'm",
            "i am",
            "my ",
            "we ",
            "our ",
            "because",
            "concern",
            "worried",
            "support",
            "oppose",
            "agree",
            "disagree",
            "insufficient",
            "enough",
            "not enough",
            "rate",
            "/10",
        )
        if any(marker in lowered for marker in opinion_markers):
            return True
        return lowered.count("?") == 0

    def _select_section_evidence(
        self,
        question: str,
        answer: str,
        evidence_pool: list[dict[str, Any]],
        *,
        limit: int = 4,
    ) -> list[dict[str, Any]]:
        if not evidence_pool:
            return []

        query_tokens = self._tokenize_evidence_text(f"{question} {answer}")
        scored: list[tuple[float, dict[str, Any]]] = []
        for item in evidence_pool:
            quote = str(item.get("quote") or "").strip()
            if not quote:
                continue

            quote_tokens = self._tokenize_evidence_text(quote)
            overlap = len(query_tokens.intersection(quote_tokens))
            base_score = float(overlap * 3)
            if str(item.get("source_type") or "") == "comment":
                base_score += 1.5

            round_no = int(item.get("round_no") or 0)
            if round_no > 0:
                base_score += min(round_no, 5) * 0.1
            base_score += min(len(quote_tokens), 40) / 20.0
            scored.append((base_score, item))

        if not scored:
            return evidence_pool[:limit]

        scored.sort(key=lambda row: row[0], reverse=True)
        selected: list[dict[str, Any]] = []
        seen_agents: set[str] = set()
        for _score, item in scored:
            if len(selected) >= limit:
                break
            agent_id = str(item.get("agent_id") or "")
            if agent_id and agent_id in seen_agents:
                continue
            selected.append(item)
            if agent_id:
                seen_agents.add(agent_id)

        if len(selected) < limit:
            existing_ids = {id(item) for item in selected}
            for _score, item in scored:
                if len(selected) >= limit:
                    break
                if id(item) in existing_ids:
                    continue
                selected.append(item)

        return selected[:limit]

    def _replace_agent_id_references(self, text: str, evidence: list[dict[str, Any]]) -> str:
        if not text:
            return text

        agent_names: dict[str, str] = {}
        for item in evidence:
            agent_id = str(item.get("agent_id") or "").strip()
            agent_name = _clean_report_text(item.get("agent_name"))
            if agent_id and agent_name:
                agent_names[agent_id.lower()] = agent_name

        if not agent_names:
            return text

        def replace_dash_id(match: re.Match[str]) -> str:
            return agent_names.get(match.group(0).lower(), match.group(0))

        rewritten = re.sub(r"\bagent-\d{1,6}\b", replace_dash_id, text, flags=re.IGNORECASE)

        def replace_numbered_agent(match: re.Match[str]) -> str:
            padded = f"agent-{int(match.group(1)):04d}"
            return agent_names.get(padded, match.group(0))

        return re.sub(r"\bagent\s*#?\s*(\d{1,6})\b", replace_numbered_agent, rewritten, flags=re.IGNORECASE)

    def _replace_post_id_references(self, text: str, evidence: list[dict[str, Any]]) -> str:
        cleaned = _clean_report_text(text)
        if not cleaned:
            return cleaned

        post_owner: dict[str, str] = {}
        fallback_name = ""
        for item in evidence:
            post_id = str(item.get("post_id") or "").strip()
            agent_name = _clean_report_text(item.get("agent_name"))
            if post_id and agent_name:
                post_owner[post_id] = agent_name
            if not fallback_name and agent_name:
                fallback_name = agent_name

        def replacement(match: re.Match[str]) -> str:
            ref_id = str(match.group(1) or "").strip()
            owner = post_owner.get(ref_id) or fallback_name
            return f"post by {owner}" if owner else "post"

        rewritten = re.sub(r"\bpost\s*id\s*#?\s*(\d+)\b", replacement, cleaned, flags=re.IGNORECASE)
        return re.sub(r"\bpost\s*#\s*(\d+)\b", replacement, rewritten, flags=re.IGNORECASE)

    def _replace_references_in_bullets(self, bullets: list[str], evidence: list[dict[str, Any]]) -> list[str]:
        rewritten: list[str] = []
        for bullet in bullets:
            text = self._replace_post_id_references(_clean_bullet_text(bullet), evidence)
            text = self._replace_agent_id_references(text, evidence)
            rewritten.append(_clean_bullet_text(text))
        return [item for item in rewritten if item]

    def _knowledge_context_lines(self, knowledge: dict[str, Any]) -> list[str]:
        if not isinstance(knowledge, dict):
            return []

        lines: list[str] = []
        summary = _clean_report_text(knowledge.get("summary"))
        if summary:
            lines.append(f"Document summary: {summary}")

        document = knowledge.get("document")
        if isinstance(document, dict):
            source_path = _clean_report_text(document.get("source_path"))
            if source_path:
                lines.append(f"Source document: {source_path}")
            text_length = document.get("text_length")
            if text_length is not None:
                lines.append(f"Document length: {text_length}")
            sources = document.get("sources")
            if isinstance(sources, list):
                for source in sources[:3]:
                    if not isinstance(source, dict):
                        continue
                    source_path = _clean_report_text(source.get("source_path"))
                    if source_path:
                        lines.append(f"Document source: {source_path}")
        return lines

    def _answer_guiding_question(
        self,
        simulation_id: str,
        question: str,
        *,
        agent_count: int,
        question_type: str,
        metric: dict[str, Any] | None,
        use_case: str | None = None,
        section_type: str = "analysis_section",
        section_title: str | None = None,
    ) -> tuple[list[str], dict[str, Any]]:
        resolved_section_title = str(section_title or question).strip() or "Report Section"
        limits = self._bullet_limits(section_type)
        knowledge_context = "\n".join(self._knowledge_context_lines(self.store.get_knowledge_artifact(simulation_id) or {}))
        style_rules = self._resolve_report_writer_instructions(use_case)
        style_block = "\n".join(f"- {rule}" for rule in style_rules)
        context_pack = self.memory.build_question_context_pack(
            simulation_id,
            question_text=question,
            metric=metric,
            question_type=question_type,
            report_title=resolved_section_title,
            limit=8,
        )
        prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "section_bullets",
            "user_template",
        )
        prompt = _render_prompt_template(
            prompt,
            simulation_id=simulation_id,
            section_kind=section_type.replace("_", " "),
            section_title=resolved_section_title,
            question=question,
            knowledge_context=knowledge_context or "- none",
            agent_count=agent_count,
            context_pack_json=json.dumps(context_pack, ensure_ascii=False),
            question_profile_json=json.dumps(context_pack.get("question_profile", {}), ensure_ascii=False),
            metric_movement_json=json.dumps(context_pack.get("metric_movement", {}), ensure_ascii=False),
            top_discourse_evidence_json=json.dumps(context_pack.get("top_discourse_evidence", []), ensure_ascii=False),
            named_agent_snippets_json=json.dumps(context_pack.get("named_agent_snippets", []), ensure_ascii=False),
            style_block=style_block or "- none",
            min_bullets=limits["min_bullets"],
            max_bullets=limits["max_bullets"],
            max_words_per_bullet=limits["max_words_per_bullet"],
        )
        return (
            self._generate_structured_bullets(
                prompt=prompt,
                failure_key=section_type,
                min_bullets=limits["min_bullets"],
                max_bullets=limits["max_bullets"],
                max_words_per_bullet=limits["max_words_per_bullet"],
            ),
            context_pack,
        )

    def _format_memory_excerpt(self, episodes: list[dict[str, Any]], limit: int = 6) -> str:
        lines: list[str] = []
        for item in episodes[: max(1, int(limit))]:
            actor_name = _clean_report_text(item.get("actor_name"))
            title = _clean_report_text(item.get("title"))
            content = _clean_report_text(item.get("content"))
            prefix = actor_name or _clean_report_text(item.get("source_kind")) or "context"
            if title:
                lines.append(f"- {prefix} | {title}: {content}")
            else:
                lines.append(f"- {prefix}: {content}")
        return "\n".join(lines)

    def _generate_structured_bullets(
        self,
        *,
        prompt: str,
        failure_key: str,
        min_bullets: int,
        max_bullets: int,
        max_words_per_bullet: int,
    ) -> list[str]:
        system_prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "section_bullets",
            "system_prompt",
        )
        repair_system_prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "section_bullets",
            "repair_system_prompt",
            default=system_prompt,
        )
        repair_template = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "section_bullets",
            "repair_user_template",
        )
        repair_attempts = self._format_repair_attempts()
        raw = ""
        try:
            raw = self.llm.complete_required(prompt, system_prompt=system_prompt)
            bullets, errors = self._parse_bullets(
                raw,
                min_bullets=min_bullets,
                max_bullets=max_bullets,
                max_words_per_bullet=max_words_per_bullet,
            )
            if not errors:
                return bullets
            for _ in range(repair_attempts):
                raw = self.llm.complete_required(
                    _render_prompt_template(
                        repair_template,
                        original_prompt=prompt,
                        original_response=raw,
                        validation_errors="\n".join(f"- {error}" for error in errors),
                        min_bullets=min_bullets,
                        max_bullets=max_bullets,
                        max_words_per_bullet=max_words_per_bullet,
                    ),
                    system_prompt=repair_system_prompt,
                )
                bullets, errors = self._parse_bullets(
                    raw,
                    min_bullets=min_bullets,
                    max_bullets=max_bullets,
                    max_words_per_bullet=max_words_per_bullet,
                )
                if not errors:
                    return bullets
        except Exception:  # noqa: BLE001
            return [self._failure_message(failure_key)]
        return [self._failure_message(failure_key)]

    def _parse_bullets(
        self,
        raw: str,
        *,
        min_bullets: int,
        max_bullets: int,
        max_words_per_bullet: int,
    ) -> tuple[list[str], list[str]]:
        try:
            payload = _parse_json_value(raw)
        except Exception:  # noqa: BLE001
            return [], ["Response must be valid JSON."]
        return _normalize_bullet_payload(
            payload,
            min_bullets=min_bullets,
            max_bullets=max_bullets,
            max_words_per_bullet=max_words_per_bullet,
        )

    def _resolve_report_writer_instructions(self, use_case: str | None) -> list[str]:
        configured: list[str] = []
        if use_case:
            try:
                configured = self.config.get_report_writer_instructions(use_case)
            except Exception:  # noqa: BLE001
                configured = []
        if configured:
            return configured
        payload = self.config.get_system_prompt_config("report_agent")
        defaults = payload.get("defaults", {}) if isinstance(payload, dict) else {}
        instructions = defaults.get("report_writer_instructions", []) if isinstance(defaults, dict) else []
        return [str(item).strip() for item in instructions if str(item).strip()]

    def _build_v2_executive_summary_from_metrics(
        self,
        *,
        simulation_id: str,
        metric_deltas: list[dict[str, Any]],
        round_count: int,
        agent_count: int,
        use_case: str | None = None,
    ) -> str:
        prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "executive_summary",
            "user_template",
        )
        prompt = _render_prompt_template(
            prompt,
            simulation_id=simulation_id,
            agent_count=agent_count,
            round_count=round_count,
            knowledge_context="\n".join(self._knowledge_context_lines(self.store.get_knowledge_artifact(simulation_id) or {})) or "- none",
            metrics_summary="; ".join(
                f"{item['metric_label']}: {item['initial_display']} -> {item['final_display']} ({item['delta']:+}{item['metric_unit']})"
                for item in metric_deltas
            ) or "- none",
            style_block="\n".join(f"- {rule}" for rule in self._resolve_report_writer_instructions(use_case)) or "- none",
        )
        system_prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "executive_summary",
            "system_prompt",
        )
        repair_system_prompt = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "executive_summary",
            "repair_system_prompt",
            default=system_prompt,
        )
        repair_template = self.config.get_system_prompt_value(
            "report_agent",
            "prompts",
            "executive_summary",
            "repair_user_template",
        )

        raw = ""
        try:
            raw = self.llm.complete_required(prompt, system_prompt=system_prompt)
            cleaned = _clean_report_text(raw)
            if cleaned and not _contains_first_person(cleaned):
                return cleaned
            for _ in range(self._format_repair_attempts()):
                raw = self.llm.complete_required(
                    _render_prompt_template(
                        repair_template,
                        original_prompt=prompt,
                        original_response=raw,
                    ),
                    system_prompt=repair_system_prompt,
                )
                cleaned = _clean_report_text(raw)
                if cleaned and not _contains_first_person(cleaned):
                    return cleaned
        except Exception:  # noqa: BLE001
            return self._failure_message("executive_summary")
        return self._failure_message("executive_summary")

    def _session_analysis_questions(self, simulation_id: str) -> list[dict[str, Any]]:
        payload = self.store.get_session_config(simulation_id)
        raw = payload.get("analysis_questions") or []
        if not isinstance(raw, list):
            return []
        return [item for item in raw if isinstance(item, dict)]

    def _session_use_case(self, simulation_id: str) -> str | None:
        use_case = str(self.store.get_session_config(simulation_id).get("use_case") or "").strip()
        return use_case or None

    def _resolve_analysis_questions(self, *, simulation_id: str, use_case: str | None) -> list[dict[str, Any]]:
        session_questions = self._session_analysis_questions(simulation_id)
        if session_questions:
            return session_questions
        if not use_case:
            return []
        try:
            return self.config.get_analysis_questions(use_case)
        except Exception:  # noqa: BLE001
            return []

    def _resolve_insight_blocks(self, use_case: str | None) -> list[dict[str, Any]]:
        if not use_case:
            return []
        try:
            return self.config.get_insight_blocks(use_case)
        except Exception:  # noqa: BLE001
            return []

    def _resolve_preset_sections(self, use_case: str | None) -> list[dict[str, Any]]:
        if not use_case:
            return []
        try:
            return self.config.get_preset_sections(use_case)
        except Exception:  # noqa: BLE001
            return []

    def _agents_from_checkpoint(self, checkpoint_records: list[dict[str, Any]]) -> list[dict[str, Any]] | None:
        if not checkpoint_records:
            return None
        rows: list[dict[str, Any]] = []
        for record in checkpoint_records:
            if not isinstance(record, dict):
                continue
            row: dict[str, Any] = {}
            metric_answers = record.get("metric_answers")
            if isinstance(metric_answers, dict):
                for metric_name, value in metric_answers.items():
                    clean_name = str(metric_name or "").strip()
                    if clean_name:
                        row[f"checkpoint_{clean_name}"] = value
            rows.append(row or dict(record))
        return rows

    def _metric_unit(self, question: dict[str, Any]) -> str:
        if question.get("type") == "yes-no":
            return "%"
        if "threshold" in question or str(question.get("metric_unit", "")).strip() == "%":
            return "%"
        return str(question.get("metric_unit", "/10") or "/10")

    def _compute_metric_value(self, question: dict[str, Any], agents: list[dict[str, Any]]) -> float:
        metric_name = str(question.get("metric_name", "")).strip()
        field = f"checkpoint_{metric_name}"
        question_type = question.get("type", "scale")
        total = max(len(agents), 1)

        if question_type == "scale":
            scores: list[float] = []
            for agent in agents:
                raw_value = agent.get(field, agent.get("opinion_post", 5.0))
                parsed = _extract_numeric_value(raw_value)
                score = parsed if parsed is not None else 5.0
                scores.append(max(0.0, min(10.0, score)))
            if "threshold" in question:
                threshold = _extract_numeric_value(question.get("threshold", 7))
                threshold_value = threshold if threshold is not None else 7.0
                return round(sum(1 for score in scores if score >= threshold_value) / total * 100, 1)
            return round(sum(scores) / len(scores) if scores else 0.0, 1)

        if question_type == "yes-no":
            yes_count = 0
            for agent in agents:
                if _parse_yes_no(agent.get(field, "")) is True:
                    yes_count += 1
            return round(yes_count / total * 100, 1)

        return 0.0

    def _bullet_limits(self, section_type: str) -> dict[str, int]:
        prompt_cfg = self.config.get_system_prompt_config("report_agent")
        defaults = prompt_cfg.get("defaults", {}) if isinstance(prompt_cfg, dict) else {}
        section_cfg = defaults.get(section_type, {}) if isinstance(defaults, dict) else {}
        return {
            "min_bullets": int(section_cfg.get("min_bullets", 3) or 3),
            "max_bullets": int(section_cfg.get("max_bullets", 5) or 5),
            "max_words_per_bullet": int(section_cfg.get("max_words_per_bullet", 32) or 32),
        }

    def _format_repair_attempts(self) -> int:
        prompt_cfg = self.config.get_system_prompt_config("report_agent")
        defaults = prompt_cfg.get("defaults", {}) if isinstance(prompt_cfg, dict) else {}
        return max(0, int(defaults.get("format_repair_attempts", 1) or 1))

    def _failure_message(self, key: str) -> str:
        prompt_cfg = self.config.get_system_prompt_config("report_agent")
        defaults = prompt_cfg.get("defaults", {}) if isinstance(prompt_cfg, dict) else {}
        failures = defaults.get("failure_messages", {}) if isinstance(defaults, dict) else {}
        message = failures.get(key)
        if message:
            return str(message).strip()
        return "Report generation failed."
